from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import os
import google.generativeai as genai
from llama_index.llms.gemini import Gemini
from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, Document
from llama_index.embeddings.gemini import GeminiEmbedding
from docx import Document as DocxDocument
import openpyxl
from dotenv import load_dotenv
import json

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("Missing GEMINI_API_KEY in .env file")

genai.configure(api_key=GEMINI_API_KEY)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

embed_model = GeminiEmbedding(model_name="models/embedding-001")
llm = Gemini(model_name="models/gemini-1.5-flash")
index = None

def extract_text_from_file(filepath):
    ext = filepath.split('.')[-1].lower()
    documents = []

    if ext == 'pdf':
        documents = SimpleDirectoryReader(input_dir=os.path.dirname(filepath)).load_data()
    elif ext == 'docx':
        doc = DocxDocument(filepath)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        documents.append(Document(text=text))
    elif ext == 'xlsx':
        workbook = openpyxl.load_workbook(filepath)
        text = ""
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            for row in worksheet.iter_rows(values_only=True):
                text += " ".join([str(cell) if cell is not None else "" for cell in row]) + "\n"
        documents.append(Document(text=text))
    elif ext == 'txt':
        with open(filepath, 'r', encoding='utf-8') as file:
            text = file.read()
        documents.append(Document(text=text))
    elif ext=='json':
        with open(filepath, 'r') as file:
            data = json.load(file)
            text = json.dumps(data, indent=4)
            # print(text)
        documents.append(Document(text=text))
    else:
        raise ValueError("Unsupported file type")

    return documents

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    allowed_extensions = {'pdf', 'docx', 'xlsx', 'txt', 'json'}
    if file and file.filename.split('.')[-1].lower() in allowed_extensions:
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            global index
            documents = extract_text_from_file(filepath)
            index = VectorStoreIndex.from_documents(documents, embed_model=embed_model)
            
            os.remove(filepath) 
            
            return jsonify({'message': 'File processed successfully'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/query', methods=['POST'])
def query():
    if not index:
        return jsonify({'error': 'Please upload a file first'}), 400
    
    query_text = request.json.get('query')
    if not query_text:
        return jsonify({'error': 'No query provided'}), 400
    
    try:
        query_engine = index.as_query_engine(llm=llm)
        response = query_engine.query(query_text)
        
        gemini_model = genai.GenerativeModel("models/gemini-1.5-flash")
        prompt = f"Based on this context: {response}, provide a clear and concise answer to: {query_text} based on the content of the uploaded file. If the information is not in the context, respond with 'I'm sorry, but I cannot answer this question based on the content of the uploaded file.'"
        gemini_response = gemini_model.generate_content(prompt)
        
        return jsonify({'response': gemini_response.text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
